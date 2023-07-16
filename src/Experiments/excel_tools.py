import pandas as pd
import os
from docx import Document
import re

## Excel-Tools contains the functions used for combining all the results of an 
## Experiment into a single Excel file as well as converting the Excel-Data
## into either multiple Word-Tables or a single Word-Table



def combine_excels(input_folder):
    """
    Combines all Excel files in the input folder into a single Excel file.
    """

    # Create a list to store the DataFrames
    data_frames = []

    # Iterate over the directory
    for root, dirs, files in os.walk(input_folder):
        for filename in files:
            # Check if the file is an Excel file
            if filename.endswith(".xlsx"):
                # Full file path
                file_path = os.path.join(root, filename)

                # Read the Excel file into a DataFrame
                data = pd.read_excel(file_path)

                # Remove all double entries
                data = data[data["Page"] != "Whole Document"]

                # Record the subdirectory name the file belongs to, e.g. ToOcr-01
                data["Subdirectory"] = os.path.basename(root)

                # Record the file name
                data["Filename"] = filename

                # Append the data DataFrame to the list
                data_frames.append(data)

    # Concatenate the DataFrames in the list
    combined_data = pd.concat(data_frames, ignore_index=True)

    # Write the combined_data DataFrame to an Excel file and save it in the Input Folder as combined_data
    output_file = os.path.join(input_folder, "combined_data.xlsx")
    combined_data.to_excel(output_file, index=False)

def excel_to_separate_wordtables(excel_doc, input_folder):
    """
    Transforms the Excel data into separate Word tables for average confidence, sum of confidence,
    runtime, and preprocessing runtime.
    """

    # Read the Excel file into a DataFrame
    excel_data = pd.read_excel(excel_doc)
    
    # Create empty Word documents for each table
    doc_average_confidence = Document()
    doc_sum_of_confidence = Document()
    doc_runtime = Document()
    doc_runtime_preprocessing = Document()

    # Extract the names of the methods used in the Experiment from the filenames
    filenames = excel_data['Filename']
    methods = set()
    pattern = r"evaluation_ocrdata_(\w+)\.xlsx"
    for filename in filenames:
        match = re.search(pattern, filename)
        if match:
            method = match.group(1)
            methods.add(method)
    
    # Create column titles for the DataFrames
    column_titles = ["DocumentNr."] + list(methods)
    
    # Create empty DataFrames for each table
    average_confidence_DataFrame = pd.DataFrame(columns=column_titles)
    sum_of_confidence_DataFrame = pd.DataFrame(columns=column_titles)
    runtime_DataFrame = pd.DataFrame(columns=column_titles)
    runtime_preprocessing_DataFrame = pd.DataFrame(columns=column_titles)
    
    # Iterate over the data grouped by subdirectory
    for toocr_document, group in excel_data.groupby("Subdirectory"):
        # Create empty dictionaries for each row
        row_average_confidence = {"DocumentNr.": toocr_document}
        row_sum_of_confidence = {"DocumentNr.": toocr_document}
        row_runtime = {"DocumentNr.": toocr_document}
        row_runtime_preprocessing = {"DocumentNr.": toocr_document}
        
        # Iterate over the methods
        for method in methods[1:]:
            # Create the filename to search for
            filename = f"evaluation_ocrdata_{method}.xlsx"
            
            try:
                # Retrieve the values from the DataFrame
                average_confidence = "{:,.2f}".format(round(group.loc[group['Filename'] == filename, 'average_confidence'].values[0], 2)).replace(",", "X").replace(".", ",").replace("X", ".")
                sum_of_confidence = round(group.loc[group['Filename'] == filename, 'sum_of_confidence'].values[0].round(), 0)
                runtime = "{:,.2f}".format(group.loc[group['Filename'] == filename, 'runtime'].values[0]).replace(",", "X").replace(".", ",").replace("X", ".")
                runtime_preprocessing = "{:,.2f}".format(group.loc[group['Filename'] == filename, 'runtime_preprocessing'].values[0]).replace(",", "X").replace(".", ",").replace("X", ".")
            except IndexError:
                # Handle missing values
                average_confidence = "NaN"
                sum_of_confidence = "NaN"
                runtime = "NaN"
                runtime_preprocessing = "Nan"
                print(f"Missing value in {toocr_document} for the method {method}")
            
            # Update the row dictionaries with the retrieved values
            row_average_confidence[method] = average_confidence
            row_sum_of_confidence[method] = sum_of_confidence
            row_runtime[method] = runtime
            row_runtime_preprocessing[method] = runtime_preprocessing
        
        # Convert the row dictionaries to DataFrames and concatenate them to the corresponding DataFrames
        row_average_confidence_df = pd.DataFrame(row_average_confidence, index=[0])
        average_confidence_DataFrame = pd.concat([average_confidence_DataFrame, row_average_confidence_df], ignore_index=True)
        
        row_sum_of_confidence_df = pd.DataFrame(row_sum_of_confidence, index=[0])
        sum_of_confidence_DataFrame = pd.concat([sum_of_confidence_DataFrame, row_sum_of_confidence_df], ignore_index=True)
        
        row_runtime_df = pd.DataFrame(row_runtime, index=[0])
        runtime_DataFrame = pd.concat([runtime_DataFrame, row_runtime_df], ignore_index=True)
        
        row_runtime_preprocessing_df = pd.DataFrame(row_runtime_preprocessing, index=[0])
        runtime_preprocessing_DataFrame = pd.concat([runtime_preprocessing_DataFrame, row_runtime_preprocessing_df], ignore_index=True)

    # Save the DataFrames as separate Word tables
    output_file_average_confidence = os.path.join(input_folder, "average_confidence.docx")
    save_table_to_word(doc_average_confidence, average_confidence_DataFrame, output_file_average_confidence)
    
    output_file_sum_of_confidence = os.path.join(input_folder, "sum_of_confidence.docx")
    save_table_to_word(doc_sum_of_confidence, sum_of_confidence_DataFrame, output_file_sum_of_confidence)

    output_file_runtime = os.path.join(input_folder, "runtime.docx")
    save_table_to_word(doc_runtime, runtime_DataFrame, output_file_runtime)

    output_file_runtime_preprocessing = os.path.join(input_folder, "runtime_preprocessing.docx")
    save_table_to_word(doc_runtime_preprocessing, runtime_preprocessing_DataFrame, output_file_runtime_preprocessing)

def excel_to_combined_wordtable(excel_doc, input_folder):
    """
    Transforms the Excel data into a combined Word table containing average confidence, 
    sum of confidence, runtime, and preprocessing runtime.
    """

    # Create a new Word document for the combined table
    doc_complete_data = Document()
    
    # Read the Excel file into a DataFrame
    excel_data = pd.read_excel(excel_doc)
    
    # Extract the names of the methods used in the Experiment from the filenames
    filenames = excel_data['Filename']
    methods = set()
    pattern = r"evaluation_ocrdata_(\w+)\.xlsx"
    for filename in filenames:
        match = re.search(pattern, filename)
        if match:
            method = match.group(1)
            methods.add(method)    
    column_titles = ["DocumentNr."] + list(methods)
    
    # Create an empty DataFrame with the column titles
    complete_DataFrame = pd.DataFrame(columns=column_titles)
    
    # Iterate over the data grouped by subdirectory
    for toocr_document, group in excel_data.groupby("Subdirectory"):
        # Create empty dictionaries for each row
        row_average_confidence = {"DocumentNr.": toocr_document}
        row_sum_of_confidence = {"DocumentNr.": toocr_document}
        row_runtime = {"DocumentNr.": toocr_document}
        row_runtime_preprocessing = {"DocumentNr.": toocr_document}
        
        # Iterate over the methods
        for method in methods:
            filename = f"evaluation_ocrdata_{method}.xlsx"
            try:
                # Retrieve the values from the DataFrame
                average_confidence = "{:,.2f}".format(round(group.loc[group['Filename'] == filename, 'average_confidence'].values[0], 2)).replace(",", "X").replace(".", ",").replace("X", ".")
                sum_of_confidence = round(group.loc[group['Filename'] == filename, 'sum_of_confidence'].values[0], 0)
                runtime = "{:,.2f}".format(group.loc[group['Filename'] == filename, 'runtime'].values[0]).replace(",", "X").replace(".", ",").replace("X", ".")
                runtime_preprocessing = "{:,.2f}".format(group.loc[group['Filename'] == filename, 'runtime_preprocessing'].values[0]).replace(",", "X").replace(".", ",").replace("X", ".")
            except IndexError:
                # Handle missing values
                average_confidence = "NaN"
                sum_of_confidence = "NaN"
                runtime = "NaN"
                runtime_preprocessing = "Nan"
                print(f"Missing value in {toocr_document} for the method {method}")
            
            # Update the row dictionaries with the retrieved values
            row_average_confidence[method] = average_confidence
            row_sum_of_confidence[method] = sum_of_confidence
            row_runtime[method] = runtime
            row_runtime_preprocessing[method] = runtime_preprocessing
        
        # Convert the row dictionaries to DataFrames and concatenate them to the complete_DataFrame
        row_sum_of_confidence_df = pd.DataFrame(row_sum_of_confidence, index=[0])
        complete_DataFrame = pd.concat([complete_DataFrame, row_sum_of_confidence_df], ignore_index=True)
        
        row_average_confidence_df = pd.DataFrame(row_average_confidence, index=[0])
        complete_DataFrame = pd.concat([complete_DataFrame, row_average_confidence_df], ignore_index=True)
        
        row_runtime_df = pd.DataFrame(row_runtime, index=[0])
        complete_DataFrame = pd.concat([complete_DataFrame, row_runtime_df], ignore_index=True)
        
        row_runtime_preprocessing_df = pd.DataFrame(row_runtime_preprocessing, index=[0])
        complete_DataFrame = pd.concat([complete_DataFrame, row_runtime_preprocessing_df], ignore_index=True)

    # Save the complete_DataFrame as a Word table
    output_file_complete = os.path.join(input_folder, "complete_data.docx")
    save_table_to_word(doc_complete_data, complete_DataFrame, output_file_complete)

def save_table_to_word(document, df, output_file):
    """
    Saves a DataFrame as a Word table in a Word document.
    """
    table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    
    # Add header
    for col in range(df.shape[1]):
        table.cell(0, col).text = str(df.columns[col])
    
    # Add data rows
    for row in range(df.shape[0]):
        for col in range(df.shape[1]):
            table.cell(row + 1, col).text = str(df.iloc[row, col])
    # Save document
    document.save(output_file)



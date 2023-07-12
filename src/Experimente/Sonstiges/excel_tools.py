import pandas as pd
import os
from docx import Document


def combine_excels(input_folder):
    # Create a list to store the DataFrames
    data_frames = []

    # Iterate over the directory
    for root, dirs, files in os.walk(input_folder):
        for filename in files:
            # Check if the file is an Excel file
            if filename.endswith(".xlsx"):
                # Full file path
                file_path = os.path.join(root, filename)

                # Read the Excel file into a DataFrame
                data = pd.read_excel(file_path)

                # Remove all double entries
                data = data[data["Page"] != "Whole Document"]

                # Record the subdirectory name the file belongs to, e.g. ToOcr-01
                data["Subdirectory"] = os.path.basename(root)

                # Record the file name
                data["Filename"] = filename

                # Append the data DataFrame to the list
                data_frames.append(data)

    # Concatenate the DataFrames in the list
    combined_data = pd.concat(data_frames, ignore_index=True)

    # Write the combined_data DataFrame to an Excel file and save it in the Input Folder as combined_data
    output_file = os.path.join(input_folder, "combined_data.xlsx")
    combined_data.to_excel(output_file, index=False)


def excel_to_wordtable(excel_doc, input_folder):
    methods = [
        "DocumentNr.",
        "original",
        "zhang",
        "medial_axis",
        "thin",
    ]
    excel_data = pd.read_excel(excel_doc)
    doc_average_confidence = Document()
    doc_sum_of_confidence = Document()
    doc_runtime = Document()
    doc_runtime_preprocessing = Document()
    
    average_confidence_DataFrame = pd.DataFrame(columns=methods)
    sum_of_confidence_DataFrame = pd.DataFrame(columns=methods)
    runtime_DataFrame = pd.DataFrame(columns=methods)
    runtime_preprocessing_DataFrame = pd.DataFrame(columns=methods)
    
    for toocr_document, group in excel_data.groupby("Subdirectory"):
        row_average_confidence = {"DocumentNr.": toocr_document}
        row_sum_of_confidence = {"DocumentNr.": toocr_document}
        row_runtime = {"DocumentNr.": toocr_document}
        row_runtime_preprocessing = {"DocumentNr.": toocr_document}
        
        for method in methods[1:]:
            filename = f"evaluation_ocrdata_{method}.xlsx"
            try:
                average_confidence = "{:,.2f}".format(round(group.loc[group['Filename'] == filename, 'average_confidence'].values[0], 2)).replace(",", "X").replace(".", ",").replace("X", ".")
                sum_of_confidence = round(group.loc[group['Filename'] == filename, 'sum_of_confidence'].values[0].round(), 0)
                runtime = "{:,.2f}".format(group.loc[group['Filename'] == filename, 'runtime'].values[0]).replace(",", "X").replace(".", ",").replace("X", ".")
                runtime_preprocessing = "{:,.2f}".format(group.loc[group['Filename'] == filename, 'runtime_preprocessing'].values[0]).replace(",", "X").replace(".", ",").replace("X", ".")
            except IndexError:
                average_confidence = "NaN"
                sum_of_confidence = "NaN"
                runtime = "NaN"
                runtime_preprocessing = "Nan"
                print(f"Missing value in {toocr_document} for the method {method}")
            
            row_average_confidence[method] = average_confidence
            row_sum_of_confidence[method] = sum_of_confidence
            row_runtime[method] = runtime
            row_runtime_preprocessing[method] = runtime_preprocessing
        
        # Convert row_average_confidence dictionary to DataFrame
        row_average_confidence_df = pd.DataFrame(row_average_confidence, index=[0])
        # Concatenate the average confidence row to the average confidence DataFrame
        average_confidence_DataFrame = pd.concat(
            [average_confidence_DataFrame, row_average_confidence_df], ignore_index=True
        )
        # Convert row_sum_of_confidence dictionary to DataFrame
        row_sum_of_confidence_df = pd.DataFrame(row_sum_of_confidence, index=[0])
        # Concatenate the sum of confidence row to the sum of confidence DataFrame
        sum_of_confidence_DataFrame = pd.concat(
            [sum_of_confidence_DataFrame, row_sum_of_confidence_df], ignore_index=True
        )
        # Convert row_runtime dictionary to DataFrame
        row_runtime_df = pd.DataFrame(row_runtime, index=[0])
        # Concatenate the runtime row to the runtime DataFrame
        runtime_DataFrame = pd.concat(
            [runtime_DataFrame, row_runtime_df], ignore_index=True
        )
        # Convert row_runtime dictionary to DataFrame
        row_runtime_preprocessing_df = pd.DataFrame(row_runtime, index=[0])
        # Concatenate the runtime row to the runtime DataFrame
        runtime_preprocessing_DataFrame = pd.concat(
            [runtime_preprocessing_DataFrame, row_runtime_preprocessing_df], ignore_index=True
        )


    # Save average confidence table
    output_file_average_confidence = os.path.join(input_folder, "average_confidence.docx")
    save_table_to_word(doc_average_confidence, average_confidence_DataFrame, output_file_average_confidence)
    
    # Save sum of confidence table
    output_file_sum_of_confidence = os.path.join(input_folder, "sum_of_confidence.docx")
    save_table_to_word(doc_sum_of_confidence, sum_of_confidence_DataFrame, output_file_sum_of_confidence)

    # Save runtime confidence table
    output_file_runtime = os.path.join(input_folder, "runtime.docx")
    save_table_to_word(doc_runtime, runtime_DataFrame, output_file_runtime)

    # Save runtime confidence table
    output_file_runtime_preprocessing = os.path.join(input_folder, "runtime_preprocessing.docx")
    save_table_to_word(doc_runtime_preprocessing, runtime_DataFrame, output_file_runtime_preprocessing)

def excel_to_wordtable_complete(excel_doc, input_folder):
    methods = [
        "DocumentNr.",
        "original",
        "zhang",
        "medial_axis",
        "thin",
    ]
    excel_data = pd.read_excel(excel_doc)
    
    doc_complete_data = Document()
    
    complete_DataFrame = pd.DataFrame(columns=methods)
    
    for toocr_document, group in excel_data.groupby("Subdirectory"):
        row_average_confidence = {"DocumentNr.": toocr_document}
        row_sum_of_confidence = {"DocumentNr.": toocr_document}
        row_runtime = {"DocumentNr.": toocr_document}
        row_runtime_preprocessing = {"DocumentNr.": toocr_document}
        
        for method in methods[1:]:
            filename = f"evaluation_ocrdata_{method}.xlsx"
            try:
                average_confidence = "{:,.2f}".format(round(group.loc[group['Filename'] == filename, 'average_confidence'].values[0], 2)).replace(",", "X").replace(".", ",").replace("X", ".")
                sum_of_confidence = round(group.loc[group['Filename'] == filename, 'sum_of_confidence'].values[0], 0)
                runtime = "{:,.2f}".format(group.loc[group['Filename'] == filename, 'runtime'].values[0]).replace(",", "X").replace(".", ",").replace("X", ".")
                runtime_preprocessing = "{:,.2f}".format(group.loc[group['Filename'] == filename, 'runtime_preprocessing'].values[0]).replace(",", "X").replace(".", ",").replace("X", ".")
            except IndexError:
                average_confidence = "NaN"
                sum_of_confidence = "NaN"
                runtime = "NaN"
                runtime_preprocessing = "Nan"
                print(f"Missing value in {toocr_document} for the method {method}")
            
            row_average_confidence[method] = average_confidence
            row_sum_of_confidence[method] = sum_of_confidence
            row_runtime[method] = runtime
            row_runtime_preprocessing[method] = runtime_preprocessing
        
        # Convert row_sum_of_confidence dictionary to DataFrame
        row_sum_of_confidence_df = pd.DataFrame(row_sum_of_confidence, index=[0])
        # Concatenate the sum of confidence row to the sum of confidence DataFrame
        complete_DataFrame = pd.concat(
            [complete_DataFrame, row_sum_of_confidence_df], ignore_index=True
        )
        # Convert row_average_confidence dictionary to DataFrame
        row_average_confidence_df = pd.DataFrame(row_average_confidence, index=[0])
        # Concatenate the average confidence row to the average confidence DataFrame
        complete_DataFrame = pd.concat(
            [complete_DataFrame, row_average_confidence_df], ignore_index=True
        )
        # Convert row_runtime dictionary to DataFrame
        row_runtime_df = pd.DataFrame(row_runtime, index=[0])
        # Concatenate the runtime row to the runtime DataFrame
        complete_DataFrame = pd.concat(
            [complete_DataFrame, row_runtime_df], ignore_index=True
        )
        # Convert row_runtime dictionary to DataFrame
        row_runtime_preprocessing_df = pd.DataFrame(row_runtime_preprocessing, index=[0])
        # Concatenate the runtime row to the runtime DataFrame
        complete_DataFrame = pd.concat(
            [complete_DataFrame, row_runtime_preprocessing_df], ignore_index=True
       )


    # Save average confidence table
    output_file_complete = os.path.join(input_folder, "complete_data.docx")
    save_table_to_word(doc_complete_data, complete_DataFrame, output_file_complete)


def save_table_to_word(document, df, output_file):
    table = document.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    
    # Add header
    for col in range(df.shape[1]):
        table.cell(0, col).text = str(df.columns[col])
    
    # Add data rows
    for row in range(df.shape[0]):
        for col in range(df.shape[1]):
            table.cell(row + 1, col).text = str(df.iloc[row, col])
    # Save document
    document.save(output_file)






combine_excels("/RIDSS2023/experiment_ergebnisse/Font/LineSpacing_Test")
#excel_to_wordtable("/RIDSS2023/experiment_ergebnisse/Kontrast/Skeletonize_Test/combined_data.xlsx", "/RIDSS2023/experiment_ergebnisse/Font/FontSize_Test")
#excel_to_wordtable_complete("/RIDSS2023/experiment_ergebnisse/Font/LineSpacing_Test/combined_data.xlsx", "/RIDSS2023/experiment_ergebnisse/Font/LineSpacing_Test")